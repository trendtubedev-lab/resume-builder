"""Render a structured tailored resume into .docx and .pdf files.

Three visual templates:
  classic  - centred header, navy section rules (default)
  banner   - full-width navy header band, left-accent section bars
  minimal  - left-aligned, charcoal accents, clean whitespace

Contact bits (email / phone / location / links) are auto-classified: emails get
a mailto: link, URLs/domains get an https: link, everything else stays plain.
Links are clickable in both PDF and DOCX.
"""
from __future__ import annotations
import io

EM = " — "   # em dash, used between title and company / degree and school
MIDDOT = "  ·  "
_LINK_TLDS = (".com", ".io", ".dev", ".net", ".org", ".me", ".co", ".ai", ".xyz")


def _g(d: dict, key: str, default=""):
    v = d.get(key, default)
    return v if v is not None else default


def _contact_bits(resume: dict) -> list:
    contact = _g(resume, "contact", {}) or {}
    bits = [contact.get("email"), contact.get("phone"), contact.get("location")]
    bits += list(contact.get("links") or [])
    return [b for b in bits if b]


def _classify_contact(bit) -> tuple:
    """Return (href, display). href is None for plain text (phone, location).

    - email          -> mailto:
    - http(s) URL    -> as-is
    - bare domain/url (linkedin.com/in/x, www.site.com, site.io) -> https://
    """
    s = str(bit).strip()
    low = s.lower()
    if low.startswith(("http://", "https://")):
        return (s, s)
    if "@" in s and "/" not in s and " " not in s:
        return ("mailto:" + s, s)
    host = low.split("/", 1)[0]
    if " " not in s and "." in s and (
        "/" in s or low.startswith("www.") or any(host.endswith(t) for t in _LINK_TLDS)
    ):
        return ("https://" + s, s)
    return (None, s)


# ── PDF ──────────────────────────────────────────────────────────────────────

def build_pdf(resume: dict, template: str = "classic") -> bytes:
    if template == "banner":
        return _pdf_banner(resume)
    if template == "minimal":
        return _pdf_minimal(resume)
    return _pdf_classic(resume)


def _esc(s):
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_contact_markup(bits, sep: str, link_color: str) -> str:
    """Build a reportlab Paragraph string with <a> links for contact bits."""
    parts = []
    for b in bits:
        href, disp = _classify_contact(b)
        d = _esc(disp)
        if href:
            parts.append(f'<a href="{_esc(href)}" color="{link_color}">{d}</a>')
        else:
            parts.append(d)
    return sep.join(parts)


def _canvas_contact(canvas, bits, x, y, font, size, base_hex, link_hex,
                    sep, centered, W) -> None:
    """Draw a contact line on a canvas, drawing link bits in `link_hex` and
    registering a clickable rectangle (canvas.linkURL) over each."""
    from reportlab.lib.colors import HexColor
    canvas.setFont(font, size)
    widths = [canvas.stringWidth(str(b), font, size) for b in bits]
    sep_w = canvas.stringWidth(sep, font, size)
    total = sum(widths) + sep_w * (len(bits) - 1)
    cur = (W / 2 - total / 2) if centered else x
    for i, b in enumerate(bits):
        if i:
            canvas.setFillColor(HexColor(base_hex))
            canvas.drawString(cur, y, sep)
            cur += sep_w
        href, disp = _classify_contact(b)
        canvas.setFillColor(HexColor(link_hex if href else base_hex))
        canvas.drawString(cur, y, disp)
        if href:
            canvas.linkURL(href, (cur, y - 1, cur + widths[i], y + size), relative=0)
        cur += widths[i]


def _role_flowable(line: str, meta: str, content_w: float, base_s, meta_hex: str):
    """Bold title on the left, right-aligned gray meta on the right.

    Rendered as a borderless 2-column table so dates sit flush to the right
    margin (matching the DOCX tab-stop behaviour) instead of running inline.
    `line`/`meta` must already be escaped.
    """
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_RIGHT
    from reportlab.lib.colors import HexColor

    title_s = ParagraphStyle("rtitle", parent=base_s, spaceBefore=0)
    cells = [Paragraph(f"<b>{line}</b>", title_s)]
    widths = [content_w]
    if meta:
        meta_s = ParagraphStyle("rmeta", parent=title_s, alignment=TA_RIGHT,
                                textColor=HexColor(meta_hex), fontSize=9)
        cells.append(Paragraph(meta, meta_s))
        widths = [content_w * 0.64, content_w * 0.36]
    t = Table([cells], colWidths=widths)
    t.setStyle(TableStyle([
        ("LEFTPADDING",   (0, 0), (-1, -1), 0),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ("VALIGN",        (0, 0), (-1, -1), "BOTTOM"),
    ]))
    return t


def _pdf_classic(resume: dict) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, HRFlowable, ListFlowable, ListItem,
    )
    W, H = LETTER
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.6*inch, bottomMargin=0.6*inch)
    cw = W - 1.5*inch
    ss = getSampleStyleSheet()
    accent = HexColor("#1A4D7A")
    name_s = ParagraphStyle("name", parent=ss["Title"], fontSize=22, alignment=TA_CENTER, spaceAfter=2)
    cont_s = ParagraphStyle("cont", parent=ss["Normal"], fontSize=9, alignment=TA_CENTER, textColor=HexColor("#444444"), spaceAfter=6)
    sec_s  = ParagraphStyle("sec",  parent=ss["Heading2"], fontSize=12, textColor=accent, spaceBefore=10, spaceAfter=2)
    body_s = ParagraphStyle("body", parent=ss["Normal"], fontSize=10.5, leading=14)
    role_s = ParagraphStyle("role", parent=body_s, fontSize=11, spaceBefore=4)

    flow = []
    def section(title):
        flow.append(Paragraph(_esc(title).upper(), sec_s))
        flow.append(HRFlowable(width="100%", thickness=0.6, color=accent, spaceAfter=4))
    def bullets(items):
        lis = [ListItem(Paragraph(_esc(b), body_s), leftIndent=10) for b in items if b]
        if lis: flow.append(ListFlowable(lis, bulletType="bullet", start="•",
                            bulletColor=HexColor("#333333"), bulletFontSize=10, leftIndent=14))
    def role(line, meta, hex_="#555555"):
        flow.append(_role_flowable(_esc(line), _esc(meta), cw, role_s, hex_))

    flow.append(Paragraph(_esc(_g(resume,"name") or "Your Name"), name_s))
    bits = _contact_bits(resume)
    if bits: flow.append(Paragraph(_pdf_contact_markup(bits, " &nbsp;|&nbsp; ", "#1A4D7A"), cont_s))
    if _g(resume,"summary"): section("Summary"); flow.append(Paragraph(_esc(resume["summary"]), body_s))
    skills = _g(resume,"skills",[]) or []
    if skills: section("Skills"); flow.append(Paragraph(", ".join(_esc(s) for s in skills), body_s))
    exp = _g(resume,"experience",[]) or []
    if exp:
        section("Experience")
        for job in exp:
            line = EM.join(x for x in [_g(job,"title"),_g(job,"company")] if x)
            m    = " | ".join(x for x in [_g(job,"location"),_g(job,"dates")] if x)
            role(line, m)
            bullets(_g(job,"bullets",[]) or [])
    edu = _g(resume,"education",[]) or []
    if edu:
        section("Education")
        for e in edu:
            line = EM.join(x for x in [_g(e,"degree"),_g(e,"school")] if x)
            m    = " | ".join(x for x in [_g(e,"location"),_g(e,"dates")] if x)
            role(line, m)
            if _g(e,"details"): flow.append(Paragraph(_esc(e["details"]), body_s))
    for block in _g(resume,"additional",[]) or []:
        hd = _g(block,"heading"); items = _g(block,"items",[]) or []
        if hd and items: section(hd); bullets(items)
    doc.build(flow)
    return buf.getvalue()


def _pdf_banner(resume: dict) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor, white
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        ListFlowable, ListItem,
    )
    W, H   = LETTER
    BAND   = 1.55 * inch
    BAR_W  = 5
    accent = HexColor("#1A3A5C")
    name_val = _g(resume,"name") or "Your Name"
    bits     = _contact_bits(resume)

    def draw_header(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(accent)
        canvas.rect(0, H-BAND, W, BAND, fill=1, stroke=0)
        canvas.setFillColor(white)
        canvas.setFont("Helvetica-Bold", 22)
        canvas.drawCentredString(W/2, H-0.72*inch, name_val)
        if bits:
            _canvas_contact(canvas, bits, 0, H-1.08*inch, "Helvetica", 9,
                            "#A8C4DC", "#FFFFFF", "  |  ", True, W)
        canvas.restoreState()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=BAND+0.2*inch, bottomMargin=0.6*inch)
    ss    = getSampleStyleSheet()
    avail = W - 1.5*inch
    sec_s  = ParagraphStyle("sec",  parent=ss["Heading2"], fontSize=11, textColor=accent, spaceBefore=0, spaceAfter=2, leftIndent=8)
    body_s = ParagraphStyle("body", parent=ss["Normal"],   fontSize=10.5, leading=14)
    role_s = ParagraphStyle("role", parent=body_s, fontSize=11, spaceBefore=4)

    flow = []
    def section(title):
        bar = Table([[Spacer(BAR_W,14), Paragraph(_esc(title).upper(), sec_s)]],
                    colWidths=[BAR_W, avail-BAR_W])
        bar.setStyle(TableStyle([
            ("BACKGROUND",    (0,0),(0,0), accent),
            ("LEFTPADDING",   (0,0),(-1,-1), 0),
            ("RIGHTPADDING",  (0,0),(-1,-1), 0),
            ("TOPPADDING",    (0,0),(-1,-1), 0),
            ("BOTTOMPADDING", (0,0),(-1,-1), 0),
            ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
        ]))
        flow.append(Spacer(1,8)); flow.append(bar); flow.append(Spacer(1,4))
    def bullets(items):
        lis = [ListItem(Paragraph(_esc(b), body_s), leftIndent=10) for b in items if b]
        if lis: flow.append(ListFlowable(lis, bulletType="bullet", start="•",
                            bulletColor=HexColor("#333333"), bulletFontSize=10, leftIndent=14))
    def role(line, meta, hex_="#555555"):
        flow.append(_role_flowable(_esc(line), _esc(meta), avail, role_s, hex_))

    if _g(resume,"summary"): section("Summary"); flow.append(Paragraph(_esc(resume["summary"]), body_s))
    skills = _g(resume,"skills",[]) or []
    if skills: section("Skills"); flow.append(Paragraph(", ".join(_esc(s) for s in skills), body_s))
    exp = _g(resume,"experience",[]) or []
    if exp:
        section("Experience")
        for job in exp:
            line = EM.join(x for x in [_g(job,"title"),_g(job,"company")] if x)
            m    = " | ".join(x for x in [_g(job,"location"),_g(job,"dates")] if x)
            role(line, m)
            bullets(_g(job,"bullets",[]) or [])
    edu = _g(resume,"education",[]) or []
    if edu:
        section("Education")
        for e in edu:
            line = EM.join(x for x in [_g(e,"degree"),_g(e,"school")] if x)
            m    = " | ".join(x for x in [_g(e,"location"),_g(e,"dates")] if x)
            role(line, m)
            if _g(e,"details"): flow.append(Paragraph(_esc(e["details"]), body_s))
    for block in _g(resume,"additional",[]) or []:
        hd = _g(block,"heading"); items = _g(block,"items",[]) or []
        if hd and items: section(hd); bullets(items)
    doc.build(flow, onFirstPage=draw_header, onLaterPages=draw_header)
    return buf.getvalue()


def _pdf_minimal(resume: dict) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, HRFlowable, ListFlowable, ListItem,
    )
    W, H     = LETTER
    HEADER_H = 1.3*inch
    name_val = _g(resume,"name") or "Your Name"
    bits     = _contact_bits(resume)

    def draw_header(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(HexColor("#2C2C2C"))
        canvas.setFont("Helvetica-Bold", 24)
        canvas.drawString(0.75*inch, H-0.65*inch, name_val)
        if bits:
            _canvas_contact(canvas, bits, 0.75*inch, H-0.95*inch, "Helvetica", 9,
                            "#777777", "#1A4D7A", MIDDOT, False, W)
        canvas.setStrokeColor(HexColor("#CCCCCC"))
        canvas.setLineWidth(0.75)
        canvas.line(0.75*inch, H-HEADER_H, W-0.75*inch, H-HEADER_H)
        canvas.restoreState()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=HEADER_H+0.15*inch, bottomMargin=0.6*inch)
    cw = W - 1.5*inch
    ss    = getSampleStyleSheet()
    sec_s  = ParagraphStyle("sec",  parent=ss["Heading2"], fontSize=10.5, textColor=HexColor("#444444"),
                             spaceBefore=12, spaceAfter=2, fontName="Helvetica-Bold")
    body_s = ParagraphStyle("body", parent=ss["Normal"],   fontSize=10.5, leading=14)
    role_s = ParagraphStyle("role", parent=body_s, fontSize=11, spaceBefore=4)

    flow = []
    def section(title):
        flow.append(Paragraph(_esc(title).upper(), sec_s))
        flow.append(HRFlowable(width="100%", thickness=0.4, color=HexColor("#AAAAAA"), spaceAfter=4))
    def bullets(items):
        lis = [ListItem(Paragraph(_esc(b), body_s), leftIndent=10) for b in items if b]
        if lis: flow.append(ListFlowable(lis, bulletType="bullet", start="•",
                            bulletColor=HexColor("#555555"), bulletFontSize=10, leftIndent=14))
    def role(line, meta, hex_="#777777"):
        flow.append(_role_flowable(_esc(line), _esc(meta), cw, role_s, hex_))

    if _g(resume,"summary"): section("Summary"); flow.append(Paragraph(_esc(resume["summary"]), body_s))
    skills = _g(resume,"skills",[]) or []
    if skills: section("Skills"); flow.append(Paragraph(", ".join(_esc(s) for s in skills), body_s))
    exp = _g(resume,"experience",[]) or []
    if exp:
        section("Experience")
        for job in exp:
            line = EM.join(x for x in [_g(job,"title"),_g(job,"company")] if x)
            m    = " | ".join(x for x in [_g(job,"location"),_g(job,"dates")] if x)
            role(line, m)
            bullets(_g(job,"bullets",[]) or [])
    edu = _g(resume,"education",[]) or []
    if edu:
        section("Education")
        for e in edu:
            line = EM.join(x for x in [_g(e,"degree"),_g(e,"school")] if x)
            m    = " | ".join(x for x in [_g(e,"location"),_g(e,"dates")] if x)
            role(line, m)
            if _g(e,"details"): flow.append(Paragraph(_esc(e["details"]), body_s))
    for block in _g(resume,"additional",[]) or []:
        hd = _g(block,"heading"); items = _g(block,"items",[]) or []
        if hd and items: section(hd); bullets(items)
    doc.build(flow, onFirstPage=draw_header, onLaterPages=draw_header)
    return buf.getvalue()


# ── DOCX ─────────────────────────────────────────────────────────────────────

def build_docx(resume: dict, template: str = "classic") -> bytes:
    if template == "banner":
        return _docx_banner(resume)
    if template == "minimal":
        return _docx_minimal(resume)
    return _docx_classic(resume)


def _set_margins(doc) -> None:
    """Match the PDF: 0.75" left/right, 0.6" top/bottom."""
    from docx.shared import Inches
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.75)
        s.top_margin = s.bottom_margin = Inches(0.6)


def _shade(para, fill_hex: str) -> None:
    """Paragraph background shading (used by the banner template)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _right_tab(p) -> None:
    """Add a right-aligned tab stop at the usable page width so a trailing
    `\\t{meta}` run pushes dates to the right margin instead of landing on the
    default 0.5" tab (where long titles collide with the date)."""
    from docx.shared import Inches
    from docx.enum.text import WD_TAB_ALIGNMENT
    # LETTER (8.5") minus 0.75" left/right margins (see _set_margins) = 7.0".
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)


def _bottom_border(para, color_hex: str, size: int = 6) -> None:
    """Add a single bottom border to a paragraph (mimics the PDF section rule).
    `size` is in eighths of a point."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _add_hyperlink(paragraph, url: str, text: str, color_hex: str, half_pt: int) -> None:
    """Append a real clickable hyperlink run to a paragraph (python-docx has no
    native API for this — build the w:hyperlink OXML with an external rel)."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), color_hex); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(half_pt)); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t); hl.append(r); paragraph._p.append(hl)


def _docx_contact(doc, bits, sep: str, size_pt: float, text_hex: str,
                  link_hex: str, align=None, shade_hex: str = None):
    """Build a contact paragraph with clickable links for emails/URLs."""
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if shade_hex:
        _shade(p, shade_hex)
    trgb = RGBColor.from_string(text_hex)
    half_pt = int(round(size_pt * 2))
    def plain(txt):
        run = p.add_run(txt); run.font.size = Pt(size_pt); run.font.color.rgb = trgb
    for i, b in enumerate(bits):
        if i:
            plain(sep)
        href, disp = _classify_contact(b)
        if href:
            _add_hyperlink(p, href, disp, link_hex, half_pt)
        else:
            plain(disp)
    return p


def _docx_body(doc, resume: dict, section_fn) -> None:
    if _g(resume,"summary"):
        section_fn("Summary"); doc.add_paragraph(resume["summary"])
    skills = _g(resume,"skills",[]) or []
    if skills:
        section_fn("Skills"); doc.add_paragraph(", ".join(str(s) for s in skills))
    exp = _g(resume,"experience",[]) or []
    if exp:
        section_fn("Experience")
        for job in exp:
            p = doc.add_paragraph()
            title = _g(job,"title"); company = _g(job,"company")
            line = f"{title}{EM}{company}" if title and company else (title or company)
            r = p.add_run(line); r.bold = True
            meta = " | ".join(x for x in [_g(job,"location"),_g(job,"dates")] if x)
            if meta: _right_tab(p); p.add_run(f"\t{meta}").italic = True
            for b in _g(job,"bullets",[]) or []: doc.add_paragraph(str(b), style="List Bullet")
    edu = _g(resume,"education",[]) or []
    if edu:
        section_fn("Education")
        for e in edu:
            p = doc.add_paragraph()
            degree = _g(e,"degree"); school = _g(e,"school")
            line = f"{degree}{EM}{school}" if degree and school else (degree or school)
            p.add_run(line).bold = True
            meta = " | ".join(x for x in [_g(e,"location"),_g(e,"dates")] if x)
            if meta: _right_tab(p); p.add_run(f"\t{meta}").italic = True
            if _g(e,"details"): doc.add_paragraph(str(e["details"]))
    for block in _g(resume,"additional",[]) or []:
        heading = _g(block,"heading"); items = _g(block,"items",[]) or []
        if heading and items:
            section_fn(heading)
            for it in items: doc.add_paragraph(str(it), style="List Bullet")


def _docx_classic(resume: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    _set_margins(doc)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(_g(resume,"name") or "Your Name"); run.bold=True; run.font.size=Pt(20)
    bits = _contact_bits(resume)
    if bits:
        _docx_contact(doc, bits, " | ", 9.5, "444444", "1A4D7A",
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    def section(title):
        p = doc.add_paragraph(); r = p.add_run(title.upper())
        r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x1A,0x4D,0x7A)
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
        _bottom_border(p, "1A4D7A")
    _docx_body(doc, resume, section)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def _docx_banner(resume: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    _set_margins(doc)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER; _shade(h,"1A3A5C")
    run = h.add_run(_g(resume,"name") or "Your Name"); run.bold=True; run.font.size=Pt(20)
    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    bits = _contact_bits(resume)
    if bits:
        _docx_contact(doc, bits, " | ", 9.5, "C5D5E8", "FFFFFF",
                      align=WD_ALIGN_PARAGRAPH.CENTER, shade_hex="2A5F8F")
    def section(title):
        p = doc.add_paragraph(); r = p.add_run(title.upper()); _shade(p,"E8EFF5")
        r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x1A,0x3A,0x5C)
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    _docx_body(doc, resume, section)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def _docx_minimal(resume: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    _set_margins(doc)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    h = doc.add_paragraph()
    run = h.add_run(_g(resume,"name") or "Your Name"); run.bold=True; run.font.size=Pt(22)
    run.font.color.rgb = RGBColor(0x2C,0x2C,0x2C)
    bits = _contact_bits(resume)
    if bits:
        _docx_contact(doc, bits, MIDDOT, 9.5, "777777", "1A4D7A")
    def section(title):
        p = doc.add_paragraph(); r = p.add_run(title.upper())
        r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x44,0x44,0x44)
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
        _bottom_border(p, "AAAAAA", size=4)
    _docx_body(doc, resume, section)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()
