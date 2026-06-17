"""Resume text extraction from PDF, DOCX, and plain text uploads.

Hardened against hostile/malformed uploads: DOCX zip-bombs are rejected before
python-docx touches them, extracted text is capped, and any unexpected parser
error is turned into a clean ValueError (the caller maps these to HTTP 400).

Note: this module is CPU-bound and may block on a crafted file. Callers MUST
run it off the event loop with a wall-clock timeout (see main.py).
"""
from __future__ import annotations

import io
import logging
import os
import zipfile

log = logging.getLogger("tailorcv.parsing")


def _int_env(name: str, default: int) -> int:
    try:
        v = int(os.getenv(name, str(default)))
        return v if v > 0 else default
    except (TypeError, ValueError):
        log.warning("Invalid %s=%r; using %d", name, os.getenv(name), default)
        return default


# Reject a DOCX whose decompressed contents exceed this, or whose compression
# ratio is implausibly high (the classic zip-bomb signature).
def _max_uncompressed() -> int:
    return _int_env("MAX_DOCX_UNCOMPRESSED_MB", 50) * 1024 * 1024


def _max_zip_ratio() -> int:
    return _int_env("MAX_DOCX_ZIP_RATIO", 100)


# Cap the text we hand downstream (to the model). A generous resume is well
# under this; the cap bounds cost/latency on huge-but-valid documents.
def _max_chars() -> int:
    return _int_env("MAX_RESUME_CHARS", 200_000)


def extract_text(filename: str, data: bytes) -> str:
    """Return plain text from an uploaded resume file.

    Supports .pdf, .docx, .txt / .md. Raises ValueError on unsupported types or
    on any file that can't be parsed safely.
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        text = _guard(filename, _from_pdf, data)
    elif name.endswith(".docx"):
        text = _guard(filename, _from_docx, data)
    elif name.endswith(".txt") or name.endswith(".md") or name.endswith(".text"):
        text = data.decode("utf-8", errors="replace")
    else:
        # Last resort: try to decode as text.
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            raise ValueError(
                f"Unsupported file type for '{filename}'. Use PDF, DOCX, or TXT."
            )
    return _cap(text)


def _guard(filename: str, fn, data: bytes) -> str:
    """Run a parser, converting any unexpected failure into a clean ValueError.

    ValueError messages are user-facing (the caller returns them as the 400
    body), so they must stay generic. Real error details are logged only.
    """
    try:
        return fn(data)
    except ValueError:
        raise  # already a clean, user-facing message
    except Exception as e:  # noqa: BLE001 - untrusted input, catch everything
        log.warning("Failed to parse %r: %s: %s", filename, type(e).__name__, e)
        raise ValueError(
            f"Could not read '{filename}'. The file may be corrupt or not a "
            "valid PDF/DOCX. Try re-exporting it, or upload a TXT version."
        )


def _cap(text: str) -> str:
    limit = _max_chars()
    if len(text) > limit:
        log.info("Truncating extracted text from %d to %d chars", len(text), limit)
        return text[:limit]
    return text


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
        # Stop early once we have enough; avoids churning through a huge PDF.
        if sum(len(p) for p in parts) > _max_chars():
            break
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError(
            "Could not extract text from this PDF (it may be a scanned image). "
            "Try exporting a text-based PDF or upload a DOCX/TXT version."
        )
    return text


def _check_docx_bomb(data: bytes) -> None:
    """Reject zip-bomb DOCX files before python-docx decompresses them."""
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile:
        raise ValueError(
            "This DOCX file is not a valid Word document. Try re-exporting it."
        )
    with zf:
        compressed = 0
        uncompressed = 0
        for info in zf.infolist():
            compressed += info.compress_size
            uncompressed += info.file_size
        max_unc = _max_uncompressed()
        if uncompressed > max_unc:
            raise ValueError(
                f"This DOCX is too large when decompressed (over "
                f"{max_unc // (1024 * 1024)} MB). It may be malformed."
            )
        # Ratio guard catches small files that explode (file_size 0 -> skip).
        if compressed > 0 and uncompressed / compressed > _max_zip_ratio():
            raise ValueError(
                "This DOCX has an abnormal compression ratio and was rejected "
                "as a potential decompression bomb."
            )


def _from_docx(data: bytes) -> str:
    _check_docx_bomb(data)

    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    # Include table cell text, which often holds resume content.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()
